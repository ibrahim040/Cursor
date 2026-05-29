#!/usr/bin/env python3
"""
Generates the revised "Work Zone TTC Device Management" synthesis as a .docx.
All inserted/changed text is rendered in BOLD so the editor can see every change.
A full "List of Changes Made" is appended at the end.
"""

from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT

doc = Document()

# ----- base styles -----
normal = doc.styles["Normal"]
normal.font.name = "Calibri"
normal.font.size = Pt(11)


def add_runs(p, text):
    """Split text on ** markers; segments inside ** are bold."""
    parts = text.split("**")
    bold = False
    for seg in parts:
        if seg:
            r = p.add_run(seg)
            r.bold = bold
        bold = not bold


def para(text="", style=None, align=None, space_after=6):
    p = doc.add_paragraph(style=style)
    if align:
        p.alignment = align
    add_runs(p, text)
    p.paragraph_format.space_after = Pt(space_after)
    return p


def h(text, level=1):
    p = doc.add_heading(level=level)
    add_runs(p, text)
    return p


def bullet(text):
    return para(text, style="List Bullet")


def numbered(text):
    return para(text, style="List Number")


def table(headers, rows, note=None, bold_all_new=False):
    """Create a table. If bold_all_new, the whole table is a new/added element."""
    t = doc.add_table(rows=1, cols=len(headers))
    t.style = "Light Grid Accent 1"
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr = t.rows[0].cells
    for i, htext in enumerate(headers):
        hdr[i].text = ""
        p = hdr[i].paragraphs[0]
        r = p.add_run(htext)
        r.bold = True
        r.font.size = Pt(9)
    for row in rows:
        cells = t.add_row().cells
        for i, val in enumerate(row):
            cells[i].text = ""
            p = cells[i].paragraphs[0]
            add_runs(p, str(val))
            for run in p.runs:
                run.font.size = Pt(9)
                if bold_all_new:
                    run.bold = True
    if note:
        para(note, space_after=10)
    return t


def figure_caption(text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_runs(p, text)
    for r in p.runs:
        r.italic = True
        r.font.size = Pt(9)
    return p


# =========================================================================
# TITLE PAGE
# =========================================================================
title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = title.add_run("Work Zone Temporary Traffic Control (TTC) Device Management")
r.bold = True
r.font.size = Pt(18)

sub = doc.add_paragraph()
sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
for line in [
    "For the Federal Highway Administration (FHWA)",
    "By Dr. Chukwuma Nnaji",
    "Department of Construction Science",
    "Texas A&M University",
]:
    rr = sub.add_run(line + "\n")
    rr.font.size = Pt(12)

note = doc.add_paragraph()
note.alignment = WD_ALIGN_PARAGRAPH.CENTER
rn = note.add_run(
    "REVISED EDITION — All text shown in BOLD indicates an addition or correction made "
    "during fact-checking and reference verification. A complete itemized List of Changes "
    "is provided at the end of this document."
)
rn.italic = True
rn.font.size = Pt(10)
rn.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

doc.add_page_break()

# =========================================================================
# 1. INTRODUCTION
# =========================================================================
h("1. Introduction", 1)

para(
    "According to data from the U.S. Bureau of Labor Statistics (BLS), the construction industry "
    "recorded **1,008** worker fatalities in 2018, **averaging nearly three deaths every day** "
    "(BLS, 2020a). The transportation and warehousing sector reported the second-highest number of "
    "workplace fatalities that year **(874 fatalities, after construction's 1,008)**. Fatality rates in "
    "the construction and transportation **and warehousing** sectors stood at 9.5 and 14.0 per 100,000 "
    "full-time equivalent workers, respectively, significantly higher than in most other industries. Only "
    "the agriculture, forestry, fishing, and hunting sector reported a higher rate, at 23.4 fatalities per "
    "100,000 workers (BLS, 2020a)."
)

para(
    "In terms of nonfatal occupational incidents, the construction industry had a total recordable injury "
    "and illness rate of 3.0 per 100 full-time workers, while the transportation and warehousing sector "
    "reported a higher rate of 4.5, both exceeding the **private-industry** average of 2.8 (BLS, 2020b). "
    "These figures correspond to approximately 200,000 reported cases in construction and around 220,000 "
    "in transportation and warehousing. A significant portion of these fatalities and injuries is linked to "
    "motor vehicle crashes, which impact both roadway workers and road users."
)

para(
    "The U.S. Department of Transportation (USDOT) reported 36,560 roadway fatalities nationwide in 2018 "
    "(USDOT, 2019), many of which occurred in or near work zones. That year, 755 work zone-related "
    "fatalities were recorded, including 124 worker deaths **(National Work Zone Safety Information "
    "Clearinghouse, 2019)**. **In 2018, large trucks or buses (commercial motor vehicles) were involved "
    "in 215 fatal work zone crashes—about 32 percent of all fatal work zone crashes that year (FHWA, "
    "2020).**"
)

para(
    "Historical data underscores the urgent need for improved safety in construction and maintenance work "
    "zones. Reducing work zone crashes remains a critical priority. Beyond the devastating emotional toll on "
    "families, coworkers, and communities, the financial burden of fatal crashes is significant, with each "
    "incident potentially costing millions of dollars (Blincoe et al., 2015). Moreover, work zone crashes "
    "have broader economic and societal impacts, including road closures, travel delays, and diminished "
    "mobility, which adversely affect state and local economies."
)

para(
    "Construction and maintenance activities conducted on active roadways typically require the "
    "implementation of Temporary Traffic Control (TTC) strategies to manage the safe and efficient flow of "
    "vehicles through or around the work zone. TTC systems are vital for safely directing motorists and "
    "protecting crews performing essential roadway work. These systems generally include signage, cones, "
    "barriers, and flaggers to alert drivers to upcoming changes and delineate safe working areas."
)

para(
    "Traffic control personnel play a central role in the deployment, maintenance, modification, and removal "
    "of TTC devices. However, the setup and takedown phases of TTC operations are particularly hazardous. "
    "During these phases, workers are often positioned close to live traffic, with limited space and "
    "visibility, placing them at significantly higher risk than during normal operations once the TTC is "
    "fully established."
)

para(
    "These transitional periods also coincide with abrupt changes in driver expectations. Motorists must "
    "quickly adjust from familiar, permanent road configurations to temporary layouts, which can lead to "
    "confusion, hesitation, or distraction, particularly when signage is unclear or traffic is congested. Such "
    "scenarios increase the risk of erratic driving behavior, sudden lane changes, or abrupt braking, "
    "heightening the danger to roadside workers."
)

para(
    "Additionally, TTC setup and removal often result in temporary bottlenecks or traffic queues, especially "
    "along high-volume corridors. These slowdowns can trigger rear-end collisions, aggressive driving, and "
    "vehicle encroachment into the work area, posing significant risks to both workers and overall traffic "
    "mobility. For these reasons, understanding and mitigating the unique hazards associated with TTC "
    "transitions is critical for ensuring comprehensive work zone safety."
)

para(
    "This synthesis and guidebook documents how agencies and contractors across different jurisdictions plan, "
    "coordinate, and execute TTC operations. It will explore key factors that influence TTC strategies, "
    "including traffic volume, roadway geometry, work zone duration, and the type of construction or "
    "maintenance activity. The review will also highlight various procedures and technologies used to move or "
    "adjust TTC devices in live traffic environments, ranging from manual handling methods to automated "
    "deployment systems and real-time traffic monitoring tools."
)

# =========================================================================
# 2. DEVICE LIFECYCLE OVERVIEW
# =========================================================================
h("2. Device Lifecycle Overview", 1)

para(
    "The lifecycle of Temporary Traffic Control Devices (TTCD) encompasses the systematic, safe, and "
    "efficient placement, adjustment, and removal of traffic control measures (cones, drums, barriers) that "
    "govern vehicle and pedestrian flow through work zones, an essential pathway directly impacting worker "
    "safety and operational success (Gambatese & Moeung, 2022; **FHWA, 2023**)."
)

para(
    "Device Lifecycle Importance for Worker Safety and Operations. Initially, every lifecycle phase involves "
    "direct worker exposure to live traffic, and incorrect timing or improper placement can result in injuries "
    "or fatalities, making lifecycle understanding critical for engineering controls, process accountability, "
    "and legal compliance. Moreover, lifecycle management supports rapid adaptation to changing field "
    "conditions, such as traffic surges, emergency vehicle access, and weather, thereby minimizing hazards and "
    "maintaining safe, navigable work zones (**FHWA, 2021**). Finally, the highest worker exposure risk occurs "
    "during TTCD placement and removal, as these require workers to operate near or within active traffic, so "
    "strict adherence to lifecycle protocols is vital (Gambatese & Moeung, 2022; ATSSA, 2008)."
)

h("a) What Is the Lifecycle of TTCD?", 2)
para(
    "The complete process of planning, deploying, modifying, and removing all TTCD spans from the initial "
    "site-specific TCP (traffic control plan) review through post-removal inspection and documentation (ATSSA, "
    "2008)."
)
para(
    "Core Phases: Pre-installation Planning covers site reconnaissance, inventory checks, hazard mapping, and "
    "worker PPE training (ATSSA, 2008). Next, Device Placement is prioritized by work zone sequence, beginning "
    "in the Advance Warning Area, then Transition, Activity, and finally Termination zones (Theiss et al., "
    "2023). Subsequently, Adjustment and Movement includes periodic staging changes, emergency adjustments, "
    "lane shifts, and mobile operations, requiring real-time monitoring and, at times, the use of Intelligent "
    "Transportation Systems (ITS) (Ullman et al., 2024). Finally, Device Removal follows the reverse of the "
    "placement sequence (Termination \u2192 Activity \u2192 Transition \u2192 Advance Warning), secures the site, "
    "restores original traffic patterns, and conducts a final hazard sweep (ATSSA, 2008)."
)
figure_caption(
    "Figure 1: Sequential Flow of TTCD Lifecycle — Placement order: Advance Warning Area \u2192 Transition "
    "Area \u2192 Activity Area \u2192 Termination Area; Removal proceeds in reverse order."
)

h("b) Device Placement, Device Types, and Sequence", 2)
para(
    "Cones, drums, barriers, portable signs, and arrow boards are deployed in strict order as per MUTCD Part 6 "
    "standards (**FHWA, 2023**; ATSSA, 2008). Order (Theiss et al., 2023) begins at the Advance Warning Area "
    "where cues for speed reduction and work zone notification are installed, then proceeds to the Transition "
    "Area where tapers are set up and traffic is realigned with spacing referenced in Table 1, continues "
    "through the Activity Area where workers and equipment are isolated and clear buffer spaces are maintained, "
    "and concludes at the Termination Area where drivers are informed that normal traffic resumes."
)

para("**Table 1: Spacing of Channelizing Devices in Tapers and Longitudinal Buffer Lengths (MUTCD, 2023; ATSSA, 2008)**")
table(
    ["Speed (mph)", "Cone/Drum Spacing in Taper (ft)", "Longitudinal Buffer (ft)"],
    [["20", "20", "115"], ["40", "40", "305"], ["60", "60", "570"], ["70", "**70**", "730"]],
)

para(
    "Worker Safety During Placement: During placement, only one worker sets each device at a time, facing "
    "oncoming traffic, and checks the surrounding area before stepping into the work zone. Workers never step "
    "outside designated buffer zones, and measurements are based on markings rather than estimates. No one "
    "works alone, as at least two workers provide visibility, redundancy, and spotter support. In addition, "
    "one worker scouts the area by scanning both ways, checking for traffic, and monitoring conditions before "
    "stepping out. Workers must not exit the buffer zone to fetch or place devices and must always remain "
    "within marked buffer limits. Devices are installed while keeping a lookout for approaching vehicles, and "
    "workers never turn their backs to oncoming traffic. All installations are performed facing the "
    "anticipated vehicle approach. Once each zone is secured, crews proceed sequentially to the next without "
    "skipping areas. In large, moving work zones, the entire set of devices moves with the active work front, "
    "adjusting advance warning, transition, and buffer sections in unison (ATSSA, 2008; Gambatese & Moeung, "
    "2022)."
)

para(
    "Noteworthy Practices from DOT Studies and Reports: Across agencies, several noteworthy practices have "
    "been documented. Minnesota DOT uses advance warning trucks and dynamic warning signs, along with mobile "
    "operations, to reduce erratic driver maneuvers during lane closures (Theiss et al., 2024). In the same "
    "vein, the Texas DOT implements convoy management strategies, such as reducing vehicle spacing, minimizing "
    "convoy length, and scheduling pullovers, to reduce unsafe passing **(Theiss et al., 2023, 2024)**. "
    "**Broader research on convoy control more generally supports the idea that structured convoy operations "
    "and clear visual indicators, such as signage, can improve coordination and driver understanding (Marjovi "
    "et al., 2015; Frydenberg et al., 2021).** Additionally, Connecticut DOT assigns specific roles for each "
    "work zone installation step and uses automated queue warning systems integrated into project workflows, "
    "leading to improved incident response and safer device adjustments (ConnDOT, 2017). Furthermore, Oregon "
    "DOT schedules device placement and removal for low-traffic periods and equips all workers with Class 2 or "
    "3 high-visibility gear per ANSI/ISEA 107-2020 (ODOT, 2011; **ANSI/ISEA, 2020**). Finally, TxDOT and "
    "others prefer semi-permanent, coordinated moving block closures using multiple shadow vehicles and "
    "truck-mounted attenuators (TMAs) for both placement and removal to shield workers in a high-speed "
    "environment (TxDOT, 2018; VDOT, 2019)."
)

para(
    "Modern Practices: Modern practices include the use of pre-built CAD templates and digital maps for layout "
    "accuracy (IDOT, 2018), as well as enhanced high-visibility apparel that must comply with ANSI/ISEA "
    "107-2020, Class 2 or 3 (ANSI/ISEA, 2020)."
)

h("c) Device Adjustment and Movement", 2)
para(
    "I. Reasons for Adjustment: Reasons for adjustment include staging changes as work phases progress, lane "
    "shifts, emergency access, or mobile operations such as rumble strip placement or roadside spraying (KDOT, "
    "2021)."
)
para(
    "II. Methods: For static zones, channels and signs are adjusted as lanes open or close. For mobile zones, "
    "devices, attenuator trucks, and shadow vehicles move with the convoy, and advance warning and "
    "end-of-zone markers are continually repositioned (PennDOT, 2020)."
)
para(
    "III. Worker Exposure Reduction: Automated and semi-automated machines for cone/barrel placement and "
    "collection reduce manual handling of cones/barrels in live traffic lanes, thereby lowering worker "
    "exposure, such as automatic cone retrieval trucks and barrel mover trucks (Gambatese & Moeung, 2022). "
    "Workers also use Truck-Mounted Attenuators (TMA) and shadow vehicles and face traffic when handling "
    "devices (Gambatese & Moeung, 2022)."
)
para(
    "IV. Noteworthy DOT Practice: Texas DOT integrates Smart Work Zones (SWZ) with ITS sensors to dynamically "
    "manage and adjust setups (**Ullman et al., 2024**)."
)

h("d) Device Removal", 2)
para(
    "I. Reverse Placement Sequence: Removal follows the reverse of placement and starts at the Termination "
    "Area, then proceeds backward through the Activity and Transition areas, and finishes at the Advance "
    "Warning Area. Workers remove traffic devices against the flow of traffic to keep protective devices in "
    "place until the end (MassDOT, 2017)."
)
para(
    "II. Worker Orientation: Workers always maintain sight of approaching traffic and, during removal, stand "
    "protected by a TMA or shadow vehicle, never in open lanes. In multi-person crews, one worker serves as a "
    "spotter, continuously monitoring errant vehicles (VDOT, 2019)."
)
para(
    "III. Automated Device Removal: Machines such as the AutoCone 500 and Epic Solutions Roadrunner are used "
    "in DOTs to automate cone pickup and reduce manual labor, referenced in (Theiss & Ullman, 2017). The "
    "advantages include dramatic reductions in worker exposure and increased efficiency in large-scale or "
    "high-speed operations (ConnDOT, 2017). The disadvantages include high capital cost and the need for "
    "specialized training and vigilant process oversight (SCDOT, 2019)."
)
para(
    "IV. Noteworthy DOT Experience: **Field evaluations of automated cone placement and retrieval equipment "
    "report reduced worker exposure to live traffic during TTCD placement and removal, which is associated "
    "with lower struck-by risk (Theiss & Ullman, 2017; Gambatese & Moeung, 2022).**"
)
para(
    "V. Post-Removal Inspection: A drive-through inspection after removal ensures all devices and debris are "
    "cleared, buffer zones are restored, and permanent signage is left unobstructed (ATSSA, 2008). "
    "Documentation of removal times, noted hazards, and restoration of clear zones is required per federal "
    "compliance (23 CFR 630, 2025)."
)
figure_caption(
    "Figure 2: Device Lifecycle Flow — Procurement & Inspection \u2192 Field Placement \u2192 Monitoring & "
    "Adjustment \u2192 Safe Removal."
)

para(
    "Noteworthy Practices from DOT Studies: MnDOT reports regular use of shadow vehicles and TMAs for all "
    "device movement and adjustment, supported by strong crew training and post-task debriefs (Work Zone "
    "Safety Consortium, 2014). TxDOT deploys Smart Work Zone ITS platforms for dynamic adjustment, real-time "
    "alerts, and traffic monitoring, paired with robust contractor guidelines (Ullman et al., 2024). Oregon "
    "DOT emphasizes protective devices, compliance with standards, contingency planning, and safe traffic "
    "control practices (**Gambatese & Moeung, 2022, 2024**)."
)

para("Table 2: Typical Spacing of Channelizing Devices (Adapted from ATSSA, 2008; **consistent with MUTCD, 2023**)")
table(
    ["Roadway Type/Speed", "Taper Spacing (ft)", "Tangent Spacing (ft)"],
    [["< 40 mph", "20", "40"], ["45 mph", "40", "80"], ["60+ mph", "60", "100"]],
)

para(
    "Use of Automated Device Placement & Removal: DOT studies and field evaluations report several automated "
    "and semi-automated systems in use for reducing direct worker exposure during TTC placement and removal "
    "(Gambatese & Moeung, 2022; Theiss & Ullman, 2017)."
)
para(
    "Fully automated systems: Caltrans Cone Machine (research prototype) is a one-operator, cab-controlled "
    "system for forward placement and bidirectional retrieval that removes the worker from the bed and "
    "roadway. Pros include eliminating lifting and platform-fall exposure and reducing struck-by risk by "
    "keeping the operator in the truck cab. Cons include that it is a research prototype only, not commercially "
    "available, and not intended for routine field deployment. Traf-Tech ACT240 (legacy commercial) is a "
    "cab-operable automated cone truck (\u2248240 cones) with programmable 25, 50, and 100-ft spacing for "
    "forward placement and forward or reverse pickup. Pros include one-person operation, consistent automated "
    "spacing, and a mature feature set for lane closures. Cons include mixed or negative field reliability in "
    "cone handling and the manufacturer's out-of-business status. Centreville AutoCone 130 (trailer) and "
    "AutoCone 500 (truck/trailer) provide high-capacity, two-sided automated placement and retrieval, with "
    "AC130 handling approximately 136 cones and AC500 approximately 500 cones, suitable for large closures. "
    "Pros include one-operator deployment, a major reduction in on-pavement and lifting exposure, and "
    "documented DOT use, for example, MnDOT's retrofit at a cost of approximately $80k. Cons include "
    "mechanical complexity, such as occasional jams, cold-weather sensitivities, such as auxiliary engine fuel "
    "or line condensation, and benefits that depend on trained, consistent operators (Theiss & Ullman, 2017)."
)
para(
    "Partially automated systems: The J-Tech Dynamic Lift System (DLS) for vertical panels is a truck-mounted "
    "hydraulic lift that lowers ergonomic loads for vertical-panel handling, not cones. Pros include a "
    "meaningful reduction in lifting strain and a fit for panel-heavy operations. Cons include that the worker "
    "remains on a moving platform, so fall exposure persists; the system is limited to panels; and adoption is "
    "gated by existing device inventories. Epic Solutions \u201cRoadrunner\u201d (CS-3100 Setter / CR-3200 "
    "Retriever) provides side-mounted setter and retriever modules for standard flatbeds handling cones up to "
    "36\u2033, with timed dispense keyed to truck speed. Pros include retrofit to existing fleets, reduced "
    "lifting, standardized spacing, positive contractor feedback after tuning, and lower capital costs than "
    "full automation. Cons include that there are typically still 1 to 2 workers on the bed, plus the driver, "
    "so platform-fall exposure remains; occasional jams occur; and it is not a true single-operator system."
)
para(
    "Automated Machine Use in Best-Practice DOT Operations: Across best-practice DOT operations, Minnesota DOT "
    "(MnDOT) demonstrates full automation with the Centreville AutoCone 500, one-operator, high-capacity, "
    "bidirectional cone placement/retrieval on a 1993 Ford F-700, used in routine lane closures for ~1 year "
    "with demos offered (Theiss & Ullman, 2017). Reported outcomes include sustained one-operator field use "
    "for ~1 year (feasibility of full automation) and a cost \u2248 $80,000 delivered in May 2016, alongside "
    "cold-weather sensitivities (aux-engine fuel lift; air-line condensation freeze-ups), occasional minor "
    "adjustments, and benefits of maintaining a small cadre of trained operators. Additionally, Delaware DOT "
    "(DelDOT) fabricated caged cone trailers with seated setters on both sides for three-person operation, "
    "later extending the cover, using brighter finishes, and adding an arrow panel. Arizona DOT (ADOT) deploys "
    "a conveyor-bed cone truck (capacity ~60\u201380 cones) with a low seat for two-person operation and "
    "on-board roll-up sign storage. California DOT (Caltrans) uses a two-person cone-truck body conceptually "
    "similar to ADOT; MnDOT also employs a front-mounted cone basket that positions the worker ahead of the "
    "truck mass for added protection (cost \u2248 $4,000, fabricated by an outside source). Missouri DOT "
    "(MoDOT) uses front- and rear-mounted baskets for 42-in two-piece (\u201ctrimline\u201d) cones, with the "
    "rear a self-standing rack for dump trucks, at costs of Front \u2248 $250 + 16 labor-hrs and Rear \u2248 "
    "$1,000 + 95 labor-hrs; a device-compatibility constraint necessitated custom baskets. Ohio DOT (ODOT) "
    "mounts a rear basket on off-season plow trucks and later added fall protection; it also operates an "
    "extended \u201cmaintenance-of-traffic\u201d truck with lower side decks/rails for a crew of three plus a "
    "driver; a fall-protection upgrade was required after an incident. Finally, the North Texas Tollway "
    "Authority (NTTA) uses a patented large-chassis TTCD truck with a tethered worker cage and hydraulic deck "
    "to raise/lower drums or cones in a ~three-worker model, with the cage costing \u2248 $20,000 plus the "
    "truck and safety equipment (Theiss & Ullman, 2017)."
)
para(
    "Limitation: In practice, automated TTC device placement/removal is constrained by curvy/complex road "
    "geometry with limited sight distance (slower driver reactions), light/glare reducing visibility, "
    "bridge/structure conditions that hinder safe device/sign positioning, budget limits on protective "
    "measures, compressed deployment/removal schedules, and low public understanding of the operation\u2019s "
    "risks (Gambatese & Moeung, 2022)."
)

# =========================================================================
# 3. PRE-INSTALLATION PLANNING AND FIELD READINESS
# =========================================================================
h("3. Pre-Installation Planning and Field Readiness", 1)
para(
    "Pre-installation planning and field readiness form the critical first layer of safety and operational "
    "efficiency in work zone traffic control, ensuring that all physical devices, personnel, and external "
    "stakeholders are primed for rapid, coordinated, and safe deployment, tailored to project-specific risks "
    "and the latest regulatory standards."
)

h("a) Inventory and Equipment Checks (with Checklist)", 2)
para(
    "A comprehensive inventory and equipment check is the foundation of safe, efficient TTC deployment, "
    "directly affecting worker protection, traffic guidance, and compliance with MUTCD and ANSI/ISEA 107-2020 "
    "standards."
)
para(
    "I. Inventory Review Prior to Site Mobilization: First, inspect all TTC devices for cleanliness, "
    "retroreflectivity, and functionality (cones, drums, signs, arrow panels, flashing lights, barriers). "
    "Next, confirm adequate quantities based on TCP-enumerated device lists (especially for long tapers, "
    "buffer spaces, and detours). In addition, high-visibility safety gear must be present and compliant with "
    "ANSI Class 2 (minimum) or Class 3 for high-speed/complex operations. Finally, review specialized "
    "equipment: truck-mounted attenuators (TMAs), shadow vehicles, portable changeable message signs (PCMS), "
    "sign supports, and automated deployment units if used (e.g., cone trucks, AutoCone 500)."
)
para("II. Pre-deployment Condition and Compliance Checklist:")
para("(Based on ATSSA Field Guide and expanded for recent DOT and NCHRP practice)")
para("Table 3: Pre-deployment Condition and Compliance Checklist")
table(
    ["Item", "OK/Needs Repair", "Quantity", "Comments"],
    [
        ["Cones, drums, channelizers \u2014 retroreflective", "", "", ""],
        ["Signs \u2014 clean, mounted at MUTCD height", "", "", ""],
        ["Portable arrow panel \u2014 charged, visible", "", "", ""],
        ["High-visibility apparel (Class 2/3)", "", "", ""],
        ["Truck-mounted attenuator \u2014 inspected", "", "", ""],
        ["PCMS \u2014 messaging program, tested", "", "", ""],
        ["First aid kit and eye wash", "", "", ""],
        ["Lighting \u2014 night operations", "", "", ""],
        ["Two-way radios/communications", "", "", ""],
        ["Fire extinguisher", "", "", ""],
        ["Worker safety gear (helmets, gloves)", "", "", ""],
    ],
)

para(
    "**The field-oriented readiness checklist below, adapted from the companion TTC Device Management Field "
    "Pocket Guide, complements Table 3 by capturing planning, coordination, and briefing tasks that should be "
    "completed before the first device is placed. The quality of field TTC is largely determined before "
    "installation begins.**"
)
para("**Table 3A: Recommended Advance-Preparation (Pre-Installation) Readiness Checklist**")
table(
    ["\u2713", "Readiness item"],
    [
        ["\u2610", "Review TCP, lane closure sequence, typical application, and special provisions."],
        ["\u2610", "Confirm work duration, roadway type, operating speed, sight distance, and access needs."],
        ["\u2610", "Inspect signs, cones, drums, barricades, supports, lights, and retroreflectivity."],
        ["\u2610", "Test arrow board, PCMS, lighting, radios, and batteries or power supply."],
        ["\u2610", "Inspect the protective vehicle and TMA; stage them before workers enter traffic exposure."],
        ["\u2610", "Confirm spotter assignments, radio channels, hand signals, and stop-work triggers."],
        ["\u2610", "Confirm high-visibility apparel and required PPE."],
        ["\u2610", "Review emergency access, escape routes, and contingency plan."],
        ["\u2610", "Notify or coordinate with law enforcement, emergency services, transit, schools, or TMC as needed."],
        ["\u2610", "Hold a crew briefing before setup."],
    ],
    note="**(Source: TTC Device Management Field Pocket Guide, Table 2; consistent with ATSSA, 2008 and FHWA, 2023.)**",
    bold_all_new=True,
)

para(
    "III. Device-Specific Example Checklist, TMA Deployment: For TMA deployment, verify that the hydraulic "
    "system is functional, the crash cushion is intact, and the lights are operational. Then confirm the "
    "arrow board is operational and set to \u201ccaution\u201d mode. Additionally, ensure spare parts and "
    "on-site repair kits are available per DOT policy, and confirm the operator is trained for both setup and "
    "emergency response (Work Zone Safety Consortium, 2014)."
)
para(
    "IV. Updating the Checklist: To keep the checklist current, incorporate digital barcoding or app-based "
    "logging of device status for large-scale operations, and add inspection fields for novel technologies "
    "such as automated cone placement systems."
)

h("b) Pre-Task Coordination with Local Agencies (Police, EMS)", 2)
para(
    "Pre-task coordination should incorporate contingency planning that outlines alternative placement methods "
    "and designated escape routes for emergencies. Clear communication with traffic control crews and planners "
    "before lane closures is essential to ensure safety and efficiency (**Gambatese & Moeung, 2022**)."
)
para(
    "Essential Coordination Steps: First, pre-job notification to local law enforcement, emergency medical "
    "services (EMS), fire departments, and local traffic management centers is required. Next, contractors "
    "should actively communicate with traffic control crews and work zone traffic control planners on safety "
    "plans and procedures that affect traffic control set-up and removal operations (Gambatese & Moeung, "
    "2022). Then, teams confirm mutual understanding of road closure timing, detour routes, and alternate "
    "access points. Finally, plan for real-time messaging by arranging for DOT or contractor notifications to "
    "local agencies (e.g., SMS or email alerts for major work zone changes or closures each day)."
)
bullet(
    "Role of Early Communication: Notifying EMS in advance ensures the fastest possible medical response to a "
    "crash, reducing time lost to roadside decision-making and misalignment between agencies, and proactive "
    "outreach also supports diversion planning, alerting local residents and businesses to allow travelers to "
    "adapt, reducing frustration and unsafe last-minute maneuvers (FHWA, 2021)."
)
bullet(
    "Public Communication Channels: Utilize variable message boards, DOT social media, and automated "
    "text/email alerts to provide local travelers with timely updates (road closed, detour active, delays "
    "expected), and engage local news outlets and community groups, especially for long-duration closures or "
    "mobile operations spanning rural corridors."
)

h("c) Worker Readiness Training", 2)
para(
    "Rigorous, documented worker training is a notably effective means of minimizing exposure to hazards, "
    "ensuring a well-drilled team can execute both deployment and emergency procedures under real-world "
    "pressures (NSC, 2024; TEEX, 2025)."
)
para(
    "I. Fundamental Training Requirements: All workers must complete a DOT- or MUTCD-compliant TTC training "
    "program tailored to their roles (installer, spotter, supervisor) (VDOT, 2025). Emphasize correct use of "
    "high-visibility ANSI/ISEA 107-2020 compliant garments, correct placement of buffer zones, and "
    "maintenance of safe escape paths at each site (ATSSA, 2008; ANSI/ISEA, 2020)."
)
para(
    "II. Specialized Modules: Training additionally covers TMA operation and safe staging, including rapid "
    "egress drills. Nighttime installation protocols focus on lighting, glare, and visible motion (NSC, 2024). "
    "Use of buffer distance, longitudinal buffer zones calculated as a function of operating speed (e.g., for "
    "55 mph, buffer space = 495 ft; see Table 4 below)."
)
para("Table 4: MUTCD-Recommended Longitudinal Buffer Distances (**MUTCD, 2023**)")
table(
    ["Operating Speed (mph)", "Buffer Distance (ft)"],
    [["20", "115"], ["30", "200"], ["40", "305"], ["55", "495"], ["65", "645"], ["70", "730"]],
)
para(
    "III. Hands-On Demonstrations and Pre-Shift Briefings: Daily hazard recognition exercises (\u201ctailgate "
    "talks\u201d), site-specific walkthroughs, and simulation drills for device installation/removal under "
    "live traffic are conducted to prime crews for field conditions. Additionally, emergency egress procedures "
    "are reinforced during toolbox safety talks, and spotter training focuses on monitoring real-time traffic "
    "approaching the scene."
)
para(
    "IV. Reference to Documentation: Training records and daily field readiness forms should be maintained by "
    "the Traffic Control Supervisor (TCS) and reviewed before each deployment."
)
figure_caption(
    "Figure 3: Pre-Installation Planning Flow — Work Zone Plan Approved \u2192 Inventory & Equipment Checks "
    "\u2192 Pre-Task Coordination with Police/EMS \u2192 Worker Safety Training Completed \u2192 Ready for TTC "
    "Installation."
)
para("Noteworthy Practices in Pre-Installation Planning:")
para(
    "I. Inventory Checklists as Live Documents: Agencies use WZ-ITS with real-time remote monitoring and "
    "time/date-stamped system data logs accessible via web dashboards; vendors provide digital copies of the "
    "logged information to project engineers (FHWA, 2021)."
)
para(
    "II. Early Agency Coordination: Additionally, some DOTs assign a liaison officer specifically to "
    "coordinate medical and emergency response for night or high-volume operations, speeding up post-crash "
    "intervention."
)
para(
    "III. Worker Training Innovations: Finally, Texas DOT requires live drills for TMA operators and spotter "
    "pairs, simulating buffer zone breaches and emphasizing TMA shadow vehicle positioning within "
    "recommended buffer distances of device installers, a method intended to reduce worker strikes."
)

# =========================================================================
# 4. TTCD PLACEMENT IN THE FIELD
# =========================================================================
h("4. TTCD Placement in the Field", 1)
para(
    "Temporary Traffic Control Device (TTCD) placement in roadwork zones is a critical element ensuring the "
    "safety of both workers and the traveling public, requiring precisely engineered layouts, context-driven "
    "strategies, and attention to human factors, especially under rapidly changing conditions."
)

h("a) Standard Work Zone Layouts: The Four Key Areas", 2)
para(
    "The Advance Warning Area begins the work zone and alerts drivers to upcoming changes in road conditions. "
    "Its placement is dictated by road speed, environment (rural or urban), and visibility, and it typically "
    "ranges from **about 350 ft in urban high-speed settings to 2,640 ft (one-half mile) on freeways and "
    "expressways**. Signs should appear sequentially: Attention, Condition, Required Action, to orient drivers "
    "to changing conditions. The Transition Area then guides drivers off their normal path, for example, by "
    "tapering a closed lane, and uses channelizing devices and arrow panels while applying calculated taper "
    "lengths in accordance with MUTCD standards (**FHWA, 2023**)."
)
figure_caption(
    "Figure 4: Component Parts of a Temporary Traffic Control Zone. **(Adapted from the MUTCD, 11th Edition, "
    "Part 6; FHWA, 2023.)**"
)
para(
    "Next, the Activity Area houses the proper workspace, including work vehicles, workers, and equipment, "
    "with longitudinal and lateral buffer zones for protection. Finally, the Termination Area allows a safe "
    "return to standard traffic conditions, concluding with downstream tapers and the removal of traffic "
    "barriers."
)

h("b) Taper Lengths and Buffer Distances", 2)
para("Table 5: Formulas for Determining Taper Length (**MUTCD, 2023**)")
table(
    ["Speed (S)", "Taper Length (L) in feet"],
    [["40 mph or less", "L = WS\u00b2 / 60"], ["45 mph or more", "L = WS"]],
)
para(
    "Where: L denotes the taper length in feet, W denotes the width of offset in feet, and S denotes the "
    "posted speed limit, the off-peak 85th-percentile speed prior to starting work, or the anticipated "
    "operating speed in mph."
)
para("Table 6: Advance Warning Sign Spacing Distances (**MUTCD, 2023; ATSSA, 2008**)")
table(
    ["Road Type", "A", "B", "C"],
    [
        ["Urban (low speed <45 mph)", "100\u2032", "100\u2032", "100\u2032"],
        ["Urban (\u226545 mph)", "350\u2032", "350\u2032", "350\u2032"],
        ["Rural", "500\u2032", "500\u2032", "500\u2032"],
        ["Freeways/Expressways", "1,000\u2032", "1,500\u2032", "2,640\u2032"],
    ],
)
para("Table 7: MUTCD Standard Taper Lengths (**MUTCD, 2023; 12-ft lane offset**)")
table(
    ["Speed (mph)", "Taper (ft; Lane Closure)"],
    [
        ["25", "250"], ["30", "300"], ["35", "350"], ["40", "400"], ["45", "450"],
        ["50", "500"], ["55", "550"], ["60", "600"], ["65", "650"],
    ],
)
para("Buffer spaces are placed between transition zones and workspaces for error recovery.")
para("Table 8: Longitudinal Buffer Space Recommendations (**MUTCD, 2023; ATSSA, 2008**)")
table(
    ["Speed (mph)", "Buffer Length (ft)"],
    [
        ["20", "115"], ["30", "200"], ["40", "305"], ["50", "425"],
        ["60", "570"], ["70", "730"], ["75", "820"],
    ],
)

para(
    "**Field Methods for Approximating Distances. When exact measurement is impractical, the following field "
    "methods support quick verification of sign spacing, taper length, and buffer space: a roller (measuring) "
    "tape for shorter distances; a vehicle odometer for longer sign-spacing distances; intermediate reference "
    "location signs where known spacing exists; survey or station markings on construction projects when "
    "visible and reliable; the skip-line method (many skip lines follow a 10-ft line with a 30-ft gap, or 40 "
    "ft from start to start); and the pacing method (the worker must know their stride length before relying "
    "on it). Measurement alone is not sufficient\u2014always perform a drive-through to confirm that drivers "
    "can see, understand, and follow the TTC sequence (TTC Device Management Field Pocket Guide).**"
)

h("c) Spacing of Channelizing Devices", 2)
para(
    "Spacing is determined by speed, taper configuration, and the nature of the work, and, per MUTCD, it "
    "typically ranges from 20\u201340 ft in tapers and 40\u201380 ft in tangent sections. For high-speed or "
    "low-visibility environments, spacing should be minimized to improve driver guidance even if setup "
    "complexity increases. In parallel, agencies are increasingly deploying automated or semi-automated "
    "machines for device placement to reduce exposure (see Table 9 below)."
)
para("Table 9: Automated and Semi-Automated Channelizing Device Placement Systems")
table(
    ["Machine Name", "System Type", "Devices/Hour", "Pros", "Cons", "Time to Deploy"],
    [
        ["AutoCone 500", "Fully Automated", "~500", "No manual exposure, rapid, high capacity", "Cost, specialized training", "<1 hr"],
        ["J-Tech Dynamic Lift", "Semi-Automated", "~200", "Reduces lifting, mid-size jobs", "Still some manual steps", "1\u20132 hrs"],
        ["Barrel Mover 5000", "Pushing System", "~100", "For barrels, remote operation", "Limited to barrels", "1 hr"],
        ["MnDOT Cone Cages", "Agency-fabricated", "~100", "Low cost, reduces manual handling", "Small-scale only", "1\u20132 hrs"],
    ],
    note="(Adapted from Gambatese & Moeung, 2022; Theiss & Ullman, 2017)",
)

para(
    "**Installing a Stationary Lane Closure \u2014 Field Step Checklist. To ensure each step occurs under "
    "increasing protection, the following sequence (adapted from the companion Field Pocket Guide) "
    "complements the spacing values above:**"
)
para("**Table 9A: Step-by-Step Installation of a Stationary Lane Closure**")
table(
    ["\u2713", "Installation step"],
    [
        ["\u2610", "Locate and mark the beginning of the workspace."],
        ["\u2610", "Measure and mark the longitudinal buffer space in advance of the work area."],
        ["\u2610", "Measure and mark the taper length based on lateral shift and operating speed."],
        ["\u2610", "Measure and mark advance warning sign locations."],
        ["\u2610", "Install advance warning signs first, beginning with the first sign drivers will encounter."],
        ["\u2610", "Place the arrow board on the shoulder before or near the beginning of the taper."],
        ["\u2610", "Install channelizing devices through the transition area with the flow of traffic."],
        ["\u2610", "Continue placing devices through the buffer and work space."],
        ["\u2610", "Install termination area devices and end-of-zone guidance."],
        ["\u2610", "Perform a drive-through inspection; document and correct deficiencies before work begins."],
    ],
    note="**Use protective vehicle or TMA protection during installation when exposure, speed, traffic volume, "
    "or agency standards warrant it. (Source: TTC Device Management Field Pocket Guide; consistent with ATSSA, "
    "2008 and FHWA, 2023.)**",
    bold_all_new=True,
)

h("d) Placement Strategies for Mobile & Short-Duration Work", 2)
para("Mobile Operations:")
para(
    "TTCDs (cones, signs, arrow boards) are typically mounted on work vehicles or shadow vehicles, and the use "
    "of truck-mounted attenuators (TMAs) and shadow vehicles is critical for worker safety. Accordingly, "
    "device deployment must enable rapid relocation, with spacing and placement adjusted for terrain and "
    "traffic gaps. In addition, in no-passing zones, advance warning signs and temporary positive barriers are "
    "recommended for worker and queue protection. **Because passing behavior on mobile operations is "
    "unpredictable, protection should not be reduced simply because the operation is moving (ATSSA, 2025; "
    "Theiss et al., 2023).**"
)
para("Short-Duration Work:")
para(
    "As noted by ATSSA (2008), setup and teardown often take longer than the actual work task, so simplified "
    "layouts and the use of high-intensity vehicle lights may substitute for cones and signs. Manual device "
    "counts can be reduced if vehicles provide adequate warning through flashing lights and high-visibility "
    "markings. Nevertheless, clear zone adherence remains vital, and devices should be placed to maximize "
    "visibility while minimizing traffic disruption. **However, simplified control must never mean "
    "unprotected: fewer devices should be offset with effective advance warning and vehicle protection, and "
    "queue risk, access conflicts, and worker exposure should be checked before opening (ATSSA, 2025; Finley "
    "et al., 2017).**"
)
para(
    "**It is useful to distinguish three duration-based categories. Short-term stationary daytime work that "
    "occupies a location for more than one hour within a single daylight period should use the applicable "
    "MUTCD recommendations, agency standard, and project TCP, with clear advance warning, a proper taper, "
    "channelization, buffer space, and safe accommodation for pedestrians, bicyclists, and access needs. "
    "Short-duration work that occupies one location for up to one hour may use simplified control when "
    "exposure is brief and site conditions are favorable. Mobile work that moves intermittently or "
    "continuously should use high-intensity lights, signs, arrow boards, and shadow vehicles where speed and "
    "traffic volume warrant (ATSSA, 2025; Finley et al., 2017).**"
)

h("e) Visibility & Human Factors", 2)
para(
    "Reflectivity and Apparel ensure that all signs and devices use retroreflective materials compliant with "
    "MUTCD and ANSI/ISEA 107-2020 for nighttime visibility, while workers are required to wear a minimum of "
    "Class 2 garments, with Class 3 preferred for greater exposure risk. In addition, nighttime installations "
    "require additional lighting on devices and personnel, and sign and device reflectivity must meet required "
    "photometric values throughout their service life."
)
para(
    "Placement on Curves & Hills (see Figure 5) follows the key principle that all critical devices and signs "
    "must be visible within the driver\u2019s stopping sight distance, so they should never be placed just "
    "beyond a crest or curve. Accordingly, advance warning devices and tapers should be positioned further "
    "upstream of curves or hills to maximize driver response time, and for crests and vertical curves, a "
    "drive-through verification of device visibility from the appropriate sight distances is recommended."
)
figure_caption(
    "Figure 5: Device Placement on Curves and Crests. **(Illustrates MUTCD sight-distance principles; FHWA, "
    "2023.)**"
)
para(
    "Noteworthy Practices for TTCD Placement: To maintain consistent driver understanding, provide "
    "sequential, unambiguous guidance at every phase using MUTCD-compliant layouts. Moreover, apply precise "
    "buffer and taper calculations based on posted speeds without compromising worker safety for short-term "
    "convenience. Likewise, in high-speed or limited-visibility environments, minimize device spacing to "
    "balance setup time and worker exposure while improving driver guidance. For mobile or short-duration "
    "work, prioritize vehicle-mounted or automated devices while ensuring minimum visibility and delineation "
    "standards under all lighting and weather conditions. Additionally, employ state-of-the-art deployment and "
    "retrieval equipment when available, supported by specialized crew training and periodic device "
    "inspections. Finally, when roadway terrain or geometry limits standard layouts, conduct field "
    "verifications such as drive-throughs to ensure all critical information is visible to drivers in time for "
    "safe maneuvering."
)

# =========================================================================
# 5. MOVEMENT AND ADJUSTMENT OF TTCD
# =========================================================================
h("5. Movement and Adjustment of TTCD", 1)
para(
    "Movement and adjustment of Temporary Traffic Control Devices (TTCD) in active work zones are complex, "
    "high-risk operations that demand meticulous planning, specialized equipment, and innovative technologies "
    "to ensure safety and operational efficiency for both workers and road users (Gambatese & Moeung, 2022)."
)

h("a) Reasons for Adjustments (Staging, Emergencies, Lane Shifts)", 2)
para(
    "Staging of Construction Phases requires that, as projects transition through planned phases, TTCDs (such "
    "as cones, barrels, and signs) must be repositioned to reflect lane shifts, changes to workspace "
    "boundaries, or the introduction of new detours, thereby maintaining safe and logical routes for "
    "motorists. Emergencies involving vehicle crashes, hazardous material spills, or severe weather events "
    "often require immediate realignment or removal of TTCDs to allow for emergency response, establish "
    "temporary detours, or secure the site, and agencies frequently hold pre-task coordination meetings with "
    "police, EMS, and state DOT for such scenarios (Gambatese & Moeung, 2022)."
)
para(
    "Lane Shifts in multi-phase projects (e.g., bridge replacement or repaving operations) may require traffic "
    "to be progressively shifted into new configurations, and each shift necessitates the movement of "
    "channelizing devices, tapers, and signage in the order dictated by the current Traffic Control Plan (TCP) "
    "and MUTCD standards (**FHWA, 2023; ATSSA, 2008**). Mobile and Short-Duration Utility Work also demands "
    "frequent adjustment for short-duration and mobile utility operations, such as raised pavement marker "
    "installations or cable repairs, which are particularly challenging because work zones are relocated "
    "repeatedly within short timeframes, resulting in increased exposure for workers (ATSSA, 2008). For "
    "example, Theiss et al. (2023) identify rumble strip installation and herbicide spraying as operations "
    "requiring continuous TTCD movement, highlighting the importance of field adaptability, and Theiss et al. "
    "(2024) similarly emphasize the need for adaptive strategies in mobile two-lane roadway operations."
)
para(
    "Finally, Agency Coordination is often required so local agencies can adopt TTCD movement strategies based "
    "on knowledge of traffic volumes, sight distances, or utility congestion (pipes, cables, etc.), which in "
    "turn demands situational judgments beyond standard plans."
)

h("b) Safe Practices in Active Traffic (Spotters, Rolling Roadblocks)", 2)
para(
    "Before any adjustment activity, teams conduct a pre-movement safety briefing covering assigned roles, "
    "communication protocols, hazard identification, and emergency plans. Next, trained spotters and flaggers "
    "maintain constant communication with workers and alert crews to oncoming traffic, erratic driver "
    "behavior, or obstructions, and their placement must preserve clear sightlines, especially on curves or in "
    "poor visibility areas (Gambatese & Moeung, 2022). For high-speed or high-volume roadways, a rolling "
    "roadblock provides a temporary buffer that allows TTCD crews to operate in a safer, traffic-free "
    "environment, as the lead vehicle, typically a TMA, proceeds at a controlled speed ahead of the work crew "
    "to slow or momentarily halt traffic. In general operations, the TMA follows the work crew, the shadow "
    "vehicle runs ahead of the traffic flow, and spotters communicate by radio. Furthermore, all workers "
    "operate behind a TMA wherever feasible, and the spacing between the shadow vehicle and the workspace "
    "depends on vehicle mass and speed (Theiss et al., 2023). During night operations, additional portable "
    "lighting and high-intensity flashing lights are required for TTCD adjustments, and law enforcement "
    "presence is recommended to slow or divert traffic (ATSSA, 2008). Finally, longitudinal buffer spaces are "
    "required by MUTCD guidance; for example, at 55 mph, a 495 ft buffer is recommended, and at 70 mph, a "
    "buffer of at least 730 ft is recommended (**MUTCD, 2023; ATSSA, 2008**)."
)
figure_caption("Figure 6: Rolling Roadblock with TMA. **(Concept after Theiss et al., 2023; Gambatese & Moeung, 2022.)**")

h("c) Equipment Used to Reduce Direct Worker Exposure", 2)
para(
    "To reduce direct worker exposure, automated and semi-automated devices play a central role. Fully "
    "automated cone machines, such as the AutoCone 500, allow cones or barrels to be deployed and collected "
    "from within the cab, drastically reducing worker exposure. In addition, dynamic lift systems (e.g., "
    "J-Tech Dynamic Lift) and cone carts\u2014manual but shielded systems such as those used by "
    "MnDOT\u2014further minimize the need for workers to enter traffic lanes, though some manual handling "
    "remains (**Theiss & Ullman, 2017**). Likewise, barrel movers and sidewinders are vehicle-mounted "
    "mechanical push-arms that enable remote repositioning of heavy drums/barrels, reducing musculoskeletal "
    "strain and direct exposure. Complementing these devices, truck-mounted attenuators (TMAs)\u2014which "
    "should now meet **AASHTO MASH (2016) crashworthiness criteria; units manufactured on or before December "
    "31, 2019, and successfully tested to NCHRP Report 350 may remain in service through their normal service "
    "lives**\u2014are essential during all movement activities and act as mobile crash cushions in case of "
    "errant vehicles. Equally important, all workers engaged in TTCD movement wear ANSI/ISEA 107-2020 "
    "compliant high-visibility garments, Class 2 at minimum for most situations and Class 3 for high-speed "
    "corridors (ANSI/ISEA, 2020). For visibility, supplemental portable lighting is mandatory during "
    "low-visibility or nighttime operations, and devices must be positioned to avoid glare while fully "
    "illuminating the activity area (ATSSA, 2008). Finally, communication technologies such as radios and "
    "hands-free systems foster rapid interaction between spotters, TMA drivers, and installers."
)

h("d) ITS-Driven Adjustments (Intelligent Transportation Systems)", 2)
para(
    "Smart Work Zones (SWZs) integrate real-time sensors (traffic detectors, cameras) and portable changeable "
    "message signs (PCMS) to monitor vehicle speed, occupancy, and queue lengths. Key functionalities include "
    "queue detection and warning, in which systems automatically update PCMS as traffic conditions change, "
    "alerting drivers to upcoming slowdowns or stopped traffic and enabling more proactive TTCD adjustment and "
    "work crew deployment (FHWA, 2021). In addition, dynamic lane merging uses traffic sensors and PCMS to "
    "guide motorists on when and where to merge based on congestion levels, thereby shaping motorist behavior "
    "and reducing unexpected lane encroachments on work crews (FHWA, 2021). Finally, variable speed limits "
    "enable automated adjustment of temporary speed limits upstream of activity areas in response to real-time "
    "congestion or incidents, thereby lowering risk during TTCD movement (FHWA, 2021)."
)
para("Benefits")
para(
    "Smart Work Zones deliver two primary benefits. First, **studies of queue-warning deployments report "
    "reductions in rear-end and secondary crashes during queuing conditions (FHWA, 2021).** Moreover, they "
    "provide actionable, context-specific data to field crews for scheduling and corridor staging, which "
    "helps agencies decide when and where to move or adjust."
)
para(
    "Note: Sensors and message signs sited in advance and through the work zone, monitored remotely for status "
    "and adjustments."
)
figure_caption(
    "Figure 7: ITS (Intelligent Transportation Systems) Integration in Work Zones — Traffic Sensors \u2192 ITS "
    "Controller \u2192 DMS/Arrow Boards \u2192 Driver Behavior Adjusts."
)
para(
    "Furthermore, ITS for Automated Lane Closures: Agencies use SWZ platforms tied to TMCs to monitor "
    "sensors/CCTV and automatically or remotely update PCMS during incidents; specifications require a "
    "wireless communication system linking detectors, controllers, and PCMS (FHWA, 2021). Likewise, Queue "
    "Warning System Calculation: Trigger thresholds are configurable; FHWA\u2019s example uses 3-minute "
    "averaging with speed bands (e.g., 40\u201345 mph for slowdown; 15\u201325 mph for stopped). The system "
    "updates PCMS; accordingly, no buffer/setup re-calculation is specified (FHWA, 2021)."
)
para(
    "Noteworthy Practices and Additional Guidance: Regarding local agency roles, local agencies must "
    "coordinate with state DOTs to identify special needs, such as school bus stops, high pedestrian "
    "crossings, or unique utility corridors requiring adapted movement strategies (e.g., compact signs, "
    "portable flagging stations), and for utility work in urban areas, agencies may permit the use of compact, "
    "movable barriers and specialized short-duration signage while imposing strict buffer requirements and "
    "requiring rapid deployment techniques; MnDOT, for instance, uses agency-fabricated cone cages and rapid "
    "deployment carts for these specialized settings. As for special utility work considerations, "
    "short-duration utility work often requires frequent TTCD movements, which pose residual risk. Specific "
    "DOTs, such as Oregon DOT and Texas DOT, have piloted rapid-deployment protocols using agency-specific "
    "guidance for short, recurring operations, such as cable pulling or valve inspections. Consequently, the "
    "principal challenges include the need for highly mobile, easily repositioned devices compatible with "
    "confined or obstructed rights-of-way (e.g., urban curb lanes or alleys)."
)
para("Table 10: TTCD Movement \u2013 Key Specifications (MUTCD, ATSSA, and Agency Reports)")
table(
    ["Parameter", "Standard/Recommendation", "Reference"],
    [
        ["Minimum Sign Height", "5 ft (rural), 7 ft (urban); clearance above grade", "**MUTCD (2023)**"],
        ["Device Spacing (Taper)", "20\u201360 ft in tapers, 40\u2013100 ft in tangent sections", "**MUTCD (2023)**"],
        ["Buffer Distance (55 mph)", "495 ft", "**MUTCD (2023); ATSSA (2008)**"],
        ["Buffer Distance (70 mph)", "730 ft", "**MUTCD (2023); ATSSA (2008)**"],
        ["Night Lighting", "Portable towers, glare-free placement", "ATSSA (2008)"],
        ["Smart Queue Activation", "Speed drops below the preset threshold for \u22652 minutes", "FHWA (2021)"],
        ["TMA Certification", "**AASHTO MASH (2016); NCHRP Report 350 units made on/before Dec 31, 2019 may remain in service**", "**AASHTO (2016)**"],
        ["Worker Apparel", "ANSI/ISEA 107-2020, Class 2 or Class 3, depending on exposure", "ANSI/ISEA (2020)"],
    ],
)

# =========================================================================
# 6. REMOVAL OF TTCD
# =========================================================================
h("6. Removal of TTCD", 1)
para(
    "Removal of TTCD is the final and most high-risk phase of work zone management, requiring the careful "
    "reversal of device placement to restore normal traffic conditions while maintaining visibility, safety "
    "margins, and worker protection. This process mitigates hazards and confusion, reduces worker exposure to "
    "live traffic, and is reinforced by post-removal inspections to confirm safe and unimpeded roadway "
    "operations (Gambatese & Moeung, 2022; ATSSA, 2008)."
)

h("a) Criteria for Safe Removal", 2)
numbered("Work Completion Confirmation: TTCD should only be removed after confirming all work zone activities are complete, and the roadway is fully prepared for open use (Gambatese & Moeung, 2022; ATSSA, 2008). Accordingly, premature removal is strictly prohibited, as it exposes workers and creates confusion among drivers.")
numbered("Worker and Equipment Clearance: All workers and mobile equipment must be outside the active travel lanes before any device is removed (ATSSA, 2008).")
numbered("Visibility Requirements: Removal should ideally be scheduled for daylight hours to maximize visibility. If conducted at night, TTCDs and crews must be supported by sufficient lighting in compliance with ANSI/ISEA 107-2020 (ATSSA, 2008).")
numbered("Protection During Removal: Truck-mounted attenuators (TMAs) and shadow vehicles must be deployed to shield workers from live traffic, especially in high-speed or multi-lane environments (Gambatese & Moeung, 2022; ATSSA, 2008).")
numbered("Reverse Order Principle: Devices must be removed in the precise reverse order of installation, starting from the termination area back to the advance warning area, to maintain clear driver guidance throughout the process (ATSSA, 2008).")
numbered("Traffic and Weather Considerations: Weather, lighting, and traffic conditions must be evaluated before initiating removal. Inclement weather or high congestion justifies postponing removal until conditions stabilize.")
numbered("Documentation and Recordkeeping: Meticulous records of device removal times, safety observations, and any deviations from the approved plan are essential for compliance and forensic review (Gambatese & Moeung, 2022).")

h("Stepwise Removal in Reverse Sequence", 2)
para(
    "The stepwise method ensures continuous driver information and protection until normal traffic conditions "
    "are fully restored. Each phase of removal follows the reverse sequence of installation, maintaining "
    "guidance and shielding workers from unnecessary exposure."
)
para("Standard Sequential Process:")
numbered("Termination Area Devices: Removal begins at the downstream end of the work zone. Crews systematically take down channelizing devices, signs, and any downstream tapers, always moving against the traffic flow to maximize protection (ATSSA, 2008).")
numbered("Activity and Buffer Area Devices: After clearing the termination zone, workers proceed to the activity area, ensuring all equipment and debris are removed. In addition, buffer zone devices are maintained until the last possible moment to preserve maximum protection, then removed systematically (ATSSA, 2008). Finally, the workspace should be cleaned and restored prior to complete device removal (ATSSA, 2008).")
numbered("Transition Area Devices: Cones and barrels are removed from the transition taper in reverse order, thereby reestablishing the prior lane configuration (ATSSA, 2008). In addition, the use of TMA-equipped vehicles positioned in front of workers is strongly advised to shield against errant vehicles (Gambatese & Moeung, 2022).")
numbered("Advance Warning Area: Removal concludes with the advance warning area, and signs must remain in place until all other devices are cleared, ensuring that drivers are informed of changing conditions until the very end. Moreover, as with other areas, signs should be removed while working against traffic flow to maximize visibility and safety (ATSSA, 2008).")

h("Post-Removal Inspection and Drive-Through", 2)
numbered("Immediate Inspection: After TTCD removal, perform a slow drive-through to spot any debris, misplaced devices, or pavement markings needing repair (ATSSA, 2008). In addition, the field supervisor should verify that all permanent regulatory and guide signs are visible, that no temporary devices remain in the clear zone, in line with AASHTO guidance (Work Zone Safety Consortium, 2014), and that traffic lanes are correctly aligned and fully open.")
numbered("Correction of Hazards: Any hazard, such as temporary markings, debris, or partial obstruction, must be remedied immediately (ATSSA, 2008). Moreover, crews should report and respond promptly to unusual driver behavior, e.g., erratic maneuvers or abrupt stops, observed during the drive-through (ATSSA, 2008).")
numbered("Permanent Restoration: Replace or reinstate permanent lane markings, repair damaged fixtures, and ensure all devices are cleared from the clear zone.")
numbered("Documentation: A log should be completed, noting date, time, deficiencies identified, corrective measures, and the crew involved. This documentation is essential for accountability, compliance, and potential forensic review (Gambatese & Moeung, 2022).")

para("Table 11: Sequential TTCD Removal Checklist")
table(
    ["Zone/Element", "Action", "Protection/Notes", "Reference"],
    [
        ["Termination Area", "Remove devices first", "Use TMA/shadow vehicle ahead of work", "ATSSA, 2008"],
        ["Activity / Buffer", "Clear devices, clean zone", "Only when the area is empty of workers and equipment", "ATSSA, 2008"],
        ["Transition Area", "Remove in reverse of the traffic flow", "Maintain channelization as long as needed", "Gambatese & Moeung, 2022"],
        ["Advance Warning Area", "Remove signs last", "Ensure driver information remains until all work ends", "ATSSA, 2008"],
        ["Entire Work Zone", "Post-removal drive-through", "Fix hazards, restore per MUTCD/agency standards", "ATSSA, 2008"],
    ],
)

para("Noteworthy Practices to Enhance TTCD Removal")
numbered("Visibility and Worker Protection: Schedule removal during daylight whenever possible. When night work is unavoidable, deploy supplemental portable lighting to maintain visibility while preventing glare. In all cases, workers are required to wear ANSI/ISEA 107-2020 Class 2 or Class 3 high-visibility safety apparel (ANSI/ISEA, 2020; ATSSA, 2008), and use shadow vehicles as protective shields, even on low-speed roads (Work Zone Safety Consortium, 2014).")
numbered("Device Removal from the Shoulder Side: Whenever feasible, remove TTCDs from the shoulder side of the work zone, reducing direct worker exposure to live traffic (ATSSA, 2008).")
numbered("Reverse Sequence Strictness: Devices must always be removed in the reverse order of placement; advance warning devices should never be removed until all devices in the termination area have been cleared.")
numbered("Adaptation for Mobile/Short Duration Work: Mobile operations should follow simplified but sequential removal; in some cases, vehicles equipped with high-intensity flashing lights may substitute for cones/signs to reduce worker exposure (ATSSA, 2008; Theiss et al., 2023).")
numbered("Integration of Automated Tools: Use automated cone pickup and removal systems (such as AutoCone 500) to minimize the need for workers to enter live travel lanes.")
numbered("Continuous Training: Workers must receive routine training focused on TTCD removal hazards, including dynamic response strategies to adapt to real-time traffic conditions.")
numbered("Documentation as Safety Assurance: Maintain comprehensive records of each removal stage, hazards encountered, and corrective actions taken. Moreover, documentation is not only a compliance requirement but also an essential tool for reinforcing safety culture and informing future improvements.")

# =========================================================================
# 7. TTCD FOR DETOUR ROUTES AND SMALL-CITY WORK ZONES
# =========================================================================
h("7. TTCD for Detour Routes and Small-City Work Zones", 1)
para(
    "The use of Temporary Traffic Control Devices (TTCD) on detour routes and in small-city work zones is "
    "essential to maintaining safe, clear, and efficient traffic flow during temporary route changes. "
    "Effective deployment requires systematic planning, precise placement, ongoing operation, and careful "
    "device removal. Their role is particularly critical in environments with limited resources, unique "
    "roadway layouts, and local constraints. When implemented properly, TTCDs reduce driver confusion, enhance "
    "worker protection, and support uninterrupted mobility throughout the affected network (ATSSA, 2008; "
    "Gambatese & Moeung, 2022)."
)

h("a) Detour Signage Setup and Takedown", 2)
para(
    "Detour planning begins with a site-specific assessment focusing on sight distance, driver expectancy, and "
    "proximity to key intersections (ATSSA, 2008; **Gambatese & Moeung, 2022**). In addition, planners must "
    "evaluate traffic volumes, adjacent land uses, and potential pedestrian or cyclist interactions before "
    "selecting detour alignments."
)
para(
    "I. Setup Sequence: Begin by installing the final detour sign first (e.g., \u201cEND DETOUR\u201d or route "
    "return sign) at the intended re-entry/termination point, ensuring drivers know where they will rejoin the "
    "original route before any closure begins. Next, install the remaining detour signs in sequence, moving "
    "backward toward the start, so that the route is fully demarcated before closing the original path (ATSSA, "
    "2008). **The detour should be opened only after the full route is signed, visible, and understandable.** "
    "In addition, signs must conform to MUTCD standards for color, retroreflectivity, and size, and placement "
    "distances should provide drivers with adequate time for safe decision-making at prevailing roadway "
    "speeds. Finally, ensure all detour signs are visible under both day and night conditions, considering "
    "portable or fixed lighting where ambient lighting is lacking."
)
para(
    "II. Takedown Sequence: Begin by removing the first sign, the initial \u201cDETOUR AHEAD\u201d or similar "
    "advance warning, and then proceed with the flow of traffic to remove upstream signs. When sufficient crew "
    "is available, use ensemble removal to minimize residual confusion; alternatively, cover signs until they "
    "are fully removed (ATSSA, 2008). **After takedown, confirm that normal route signs and pavement markings "
    "are visible and not conflicting (TTC Device Management Field Pocket Guide).**"
)
figure_caption("Figure 8: TTCD in a Small-City Detour.")
figure_caption(
    "Figure 9: Detour Signage Deployment Diagram — Road Closed (Urban Intersection) \u2192 Detour Sign w/ "
    "Arrow (Utility Work Area) \u2192 Confirmation Sign \u2192 Flagger Zone."
)
para(
    "Detour management must explicitly account for pedestrians and bicyclists whenever sidewalks, shared-use "
    "paths, shoulders, or bike lanes are impacted. At a minimum, the TTC plan should (1) identify whether "
    "pedestrian/bicycle travel will be maintained through the work zone or rerouted via a signed detour, (2) "
    "provide a continuous, accessible route (e.g., ADA-compliant pedestrian path with safe crossings where "
    "needed), and (3) use appropriate channelization/signing to keep users separated from live traffic and "
    "work activity. The plan should also ensure that TTCDs do not create new hazards in the pedestrian/bicycle "
    "clear paths."
)

h("b) Rural and Urban Differences in TTCD Strategies", 2)
para(
    "I. Rural Contexts: Rural detours may involve longer routes with fewer alternate access points, and sign "
    "spacing is often greater due to higher travel speeds and reduced lighting. Devices are spaced according "
    "to MUTCD rural guidelines, with a typical 500 ft of advance warning, which may increase to "
    "1,500\u20132,640 ft for high-speed arterials. Buffer and taper lengths are extended to accommodate higher "
    "approach speeds and longer stopping distances. Reliance on static signage is higher, and portable "
    "changeable message signs may be used if resources allow."
)
para(
    "II. Urban/Small-City Contexts: In urban or small-city settings, shorter sign spacings, frequent "
    "intersections, and denser pedestrian/bike environments require tailored device layouts and may "
    "necessitate sidewalk detour signage. Accordingly, channelizing devices (cones/barrels) are placed at "
    "tighter intervals, often 20\u201360 ft in tapers, and additional wayfinding signage is used to address "
    "the numerous local destinations. Effective implementation also demands stakeholder coordination with "
    "transit agencies, schools, and emergency services, while placing greater emphasis on advanced public "
    "communication and traveler information through local media, social networks, and real-time notification "
    "systems."
)
para(
    "**Table 11A: Decisions for Small-City Contexts (adapted from the TTC Device Management Field Pocket "
    "Guide). Small-city work zones often operate within short blocks, closely spaced intersections, "
    "driveways, on-street parking, schools, transit stops, and active pedestrian and bicycle routes, so TTC "
    "layouts should be compact, easy to understand, and sensitive to local access needs while still giving "
    "drivers enough time to see, understand, and respond.**"
)
table(
    ["Small-city condition", "Recommended TTC approach", "Key caution"],
    [
        ["Short blocks and frequent intersections", "Use shorter, site-specific sign spacing based on speed, block length, sight distance, and nearby decision points.", "Avoid placing too many signs too close together; messages must remain clear and actionable."],
        ["Pedestrian and bicycle activity", "Provide sidewalk detour signs, pedestrian channelization, bike-route guidance, and accessible paths where needed.", "Do not route pedestrians or bicyclists into live traffic without clear protection."],
        ["Constrained streets or limited shoulders", "Use compact, portable devices and tighter channelization where appropriate for curb lanes, alleys, and narrow streets.", "Maintain a usable travel path, worker escape space, and visibility around parked vehicles or buildings."],
        ["Local access and community impacts", "Coordinate with schools, transit agencies, emergency services, and local stakeholders before closures or detours.", "Small detours can create major local disruption; provide advance public notice when possible."],
    ],
    bold_all_new=True,
)

h("c) Scaled Layouts for Low-Budget Local Agencies", 2)
numbered("Standardized modular layouts (template-based) facilitate effective traffic control in small organizations with limited resources.")
numbered("Device sharing or pooling across local agencies in neighboring jurisdictions maximizes resource efficiency.")
numbered("Scaled-down setups prioritize: critical advance warning signs, at minimum \u201cROAD WORK AHEAD\u201d and arrow boards or detour path arrows as feasible, while maintaining minimal channelization using cones/drums at proper spacing with a focus on approaches and critical decision points, and maximizing the use of existing infrastructure (e.g., streetlights for illumination, permanent posts for temporary signage).")
numbered("Innovative local cost-saving examples: multi-purpose signs with interchangeable panels, renting or borrowing portable PCMS for the project duration, and the use of agency-fabricated cone carts or manual placement devices for efficient setup and takedown.")
para("Table 12: Example Scaled Layouts for Local Agencies")
table(
    ["Layout Type", "Min. Warning Signs", "Channelizing Devices", "Portable Lighting", "PCMS Use"],
    [
        ["Low-budget, Rural", "2", "8\u201312 (per approach)", "As resources allow", "Seldom, only if available"],
        ["Small Urban", "3\u20134", "12\u201320 (per approach)", "Use streetlights", "Rented or agency-owned"],
    ],
)

h("d) Flagging in Small-Town or Low-Speed Environments", 2)
para(
    "Flaggers remain the most flexible and effective strategy in small cities and low-speed detours, "
    "particularly at intersections or alternating single-lane routes (ATSSA, 2008; Theiss et al., 2023)."
)
para(
    "Noteworthy Practices: Use Class 2 or 3 ANSI/ISEA 107-2020 high-visibility apparel and carry "
    "LED-illuminated paddles for low-light conditions, and position flaggers at locations affording maximum "
    "approach sight distance with viable retreat options. When distance, visibility, or winding roads "
    "necessitate more than one flagger, employ two-way radios or visual signals for coordinated control. For "
    "extremely short detours or alternating lane control, deploy Automated Flagger Assistance Devices (AFADs) "
    "or \u201cSTOP/SLOW\u201d paddles mounted on tripods when flaggers are unavailable (Theiss et al., 2023). "
    "In no-shoulder or narrow-location situations, protect flaggers with TMAs or work vehicles whenever "
    "possible."
)
para("Table 13: Summary of Flagging Practice Contexts")
table(
    ["Environment", "Typical Max. Speed", "Noteworthy Devices", "Coordination"],
    [
        ["Rural 2-lane", "55 mph", "Roadside flagger, TMA", "Hand/radio/sight cue"],
        ["Small city", "25\u201335 mph", "Two flaggers, sign paddles", "Two-way radios"],
        ["Urban detour", "25\u201340 mph", "Multiple flaggers/AFADs", "Central supervisor"],
    ],
)

h("e) Temporary Traffic Control for Utility Work", 2)
para(
    "Utility work in detour or small-city settings requires rapid, frequently short-term deployment of traffic "
    "control, often with limited advance public notification (ATSSA, 2008; **Utility Work Zone Traffic Control "
    "Guidelines, 2008**)."
)
para(
    "Key elements: Use portable, flexible signage, cones, and barriers for fast deployment and removal. "
    "Moreover, for single-lane closures, rolling setups (shadow vehicle with signs/arrows and worker vehicles "
    "in convoy) are recommended for maximum safety and mobility (Theiss et al., 2023). Next, the advance "
    "warning sign minimum is \u201cUTILITY WORK AHEAD\u201d placed at least 200\u2013500 ft in advance, with "
    "additional signage at decision or merge points as dictated by road speed (ATSSA, 2008). Where permissible "
    "under local standards, use a \u201cwork vehicle with flashing beacon\u201d in lieu of a full sign array "
    "for very short-duration work (<1 hour) (ATSSA, 2008). Additionally, all utility workers in or adjacent to "
    "the right-of-way must wear appropriate high-visibility PPE at all times, and for longer-duration utility "
    "work, portable rumble strips or temporary bump strips may be added to increase driver alertness "
    "(Gambatese & Moeung, 2022)."
)
para(
    "**Because utility operations often occur in constrained rights-of-way\u2014urban curb lanes, alleys, or "
    "small-city streets\u2014TTC should be portable, flexible, and easy to reposition while still protecting "
    "workers and road users. Repeated TTCD movement increases exposure, so device position and visibility "
    "should be rechecked after each move. The table below summarizes recommended approaches by utility-work "
    "condition (adapted from the TTC Device Management Field Pocket Guide).**"
)
para("**Table 13A: Utility Work TTC Decisions**")
table(
    ["Utility work condition", "Recommended TTC approach", "Key caution"],
    [
        ["Short-duration work: brief inspection, repair, or access work.", "Use portable signs, cones, and work-vehicle warning lights for quick setup and removal.", "Verify speed, sight distance, traffic volume, and worker exposure before reducing devices."],
        ["Lane closure or narrowed travel path: work blocks a lane, curb lane, or shoulder.", "Use advance warning, channelization, an arrow board, and a shadow vehicle where speed and volume warrant. Place \u201cUTILITY WORK AHEAD\u201d or equivalent signs based on agency spacing.", "Provide adequate taper and buffer space; avoid exposing workers to late merges or constrained escape paths."],
        ["Mobile or recurring work: work moves along a corridor or requires repeated stops.", "Use a rolling/convoy setup with a shadow vehicle, signs or arrows, and movable devices.", "Repeated TTCD movement increases exposure; recheck device position and visibility after each move."],
        ["Longer-duration or higher-risk work: limited sight distance, queues, or high-speed traffic.", "Use a fuller TTC setup with advance warning, channelization, buffer space, and additional alerting devices such as rumble strips where appropriate.", "Inspect and adjust TTC as lighting, weather, traffic, or underground conditions change."],
    ],
    note="**(Source: TTC Device Management Field Pocket Guide, Table 6; Utility Work Zone Traffic Control Guidelines, 2008; ATSSA, 2008.)**",
    bold_all_new=True,
)
figure_caption(
    "Figure 10: Sample Detour Layout with Optional Flagging and Utility Operation Setup (MUTCD)."
)
para("Noteworthy Practices and Challenging Moments:")
para(
    "Maintaining visibility requires frequent inspections and adjustments due to rapidly changing lighting, "
    "weather, or work progress, helping prevent driver confusion and worker injury (ATSSA, 2008). Likewise, "
    "documentation mandates that every device deployment, movement, or removal operation be logged, including "
    "time, personnel, equipment, and any deviations from the approved Traffic Control Plan (Gambatese & "
    "Moeung, 2022). In parallel, community communication entails notifying local stakeholders (emergency "
    "services, transit authorities, schools) well in advance of all detours and significant utility "
    "operations, recognizing the outsized impact on small-city mobility systems. Furthermore, adapting to "
    "constraints means small agencies must innovate, combining legacy methods (manual cone placement, "
    "hand-painted detour guides) with new technologies (shared PCMS, drone-based monitoring, GPS-based layout "
    "checks). Consistently, worker safety discipline emphasizes the use of truck-mounted attenuators (TMA) and "
    "shadow vehicles in even the smallest projects, which has proven effective at reducing struck-by "
    "incidents, especially during transition and takedown phases (Gambatese & Moeung, 2022; Theiss et al., "
    "2023). Finally, rapid responses through real-time adjustments, enabled by good communication and mobile "
    "TTCDs, are vital in small cities, where rerouting traffic unexpectedly due to utility emergencies or "
    "shifting construction needs is frequent."
)

# =========================================================================
# 8. WORKER SAFETY PROTOCOLS
# =========================================================================
h("8. Worker Safety Protocols During TTC Activities", 1)
para(
    "Worker safety protocols during Temporary Traffic Control (TTC) activities provide structured measures, "
    "such as the mandatory use of personal protective equipment (PPE), deployment of protective/impact "
    "vehicles, adherence to safety communication protocols, and proactive hazard identification, to safeguard "
    "workers in high-risk roadway environments. These protocols are especially critical during the placement, "
    "adjustment, and removal of TTC devices, which represent the most hazardous operations in work zones due "
    "to direct worker exposure to live traffic. Their implementation reduces struck-by incidents, enhances "
    "situational awareness, and ensures compliance with federal and state safety standards, thereby preventing "
    "accidents and fatalities (ATSSA, 2008)."
)

h("a) High-Visibility Safety Apparel: ANSI/ISEA 107-2020 Gear Guidelines", 2)
para(
    "I. Mandatory Apparel Requirements: All workers on the highway right-of-way must wear high-visibility "
    "safety apparel compliant with ANSI/ISEA 107-2020, and the minimum requirement is Class 2 for general TTC "
    "activities, while Class 3 (including sleeves and additional retroreflective material) is specified for "
    "night work, high-speed roads, or any situation involving greater worker exposure."
)
para(
    "II. Performance Standards: Garments must provide conspicuity under varied lighting, including dawn, dusk, "
    "and darkness under vehicle headlights. Accordingly, strict requirements are set for background color, "
    "amount of retroreflective material, garment design, placement, and durability."
)
para(
    "III. Garment Durability and Markings: Garments must retain photometric and physical performance after "
    "repeated laundering, and labels must indicate care instructions and compliance. Additionally, guidelines "
    "address single-use coveralls for environments causing rapid soiling."
)
para("Table 14: ANSI/ISEA 107-2020 Performance Classes and Typical Applications")
table(
    ["Class", "Typical Use", "Features", "Minimum Material"],
    [
        ["2", "General TTC, daylight, moderate speed", "Vests with retroreflective bands on the torso, over the shoulders", "775 in\u00b2 background, 201 in\u00b2 reflective"],
        ["3", "High-speed, complex zones, nighttime", "Includes sleeves, a greater retroreflective surface, and may add pants", "1,240 in\u00b2 background, 310 in\u00b2 reflective"],
    ],
)

h("b) Truck-Mounted Attenuator (TMA) Deployment", 2)
para(
    "I. Device Description and Compliance: TMAs are vehicle-mounted devices designed and crash-tested for "
    "impact attenuation, protecting workers from errant vehicles. **The current crashworthiness standard is "
    "AASHTO MASH (2016); TMAs manufactured after December 31, 2019, must have been successfully tested to "
    "MASH, while units manufactured on or before that date and successfully tested to NCHRP Report 350 may "
    "remain in service throughout their normal service lives (AASHTO, 2016).** Accordingly, deployment is "
    "mandatory during mobile/removal operations on high-speed or high-volume roadways or whenever worker "
    "exposure is elevated (Gambatese & Moeung, 2022). Likewise, the use of shadow vehicles equipped with TMAs "
    "is essential during the installation, movement, and removal of cones, barrels, or other TTC devices "
    "(ATSSA, 2008)."
)
para(
    "II. Spatial Positioning and Buffer Space: The TMA must be positioned in advance of the active work area, "
    "maintaining proper longitudinal and lateral buffer distances as recommended by MUTCD/State agency "
    "guidelines (e.g., minimum roll-ahead distance for the given speed/class of TMA). For mobile operations, "
    "the shadow vehicle follows the work vehicle at a buffer space sufficient to absorb energy and prevent "
    "intrusion into the work area (ATSSA, 2008)."
)
para(
    "III. Operational Noteworthy Practices: TMA-equipped vehicles should have high-intensity rotating, "
    "flashing, or strobe lights and, where practical, mounted arrow boards to increase conspicuity. Moreover, "
    "TMAs must be inspected routinely, and damage or improper positioning warrant immediate correction or "
    "withdrawal from service (Gambatese & Moeung, 2022)."
)

h("c. Safety Briefings, Toolbox Talks, and Hazard Recognition", 2)
para(
    "I. Daily Pre-Shift Briefings: Every shift begins with a comprehensive safety briefing, commonly termed "
    "tailgate, toolbox, or \u201cdeep\u201d briefings, and these meetings outline specific tasks, review known "
    "and anticipated hazards, assign roles, and review TTC layout, emergency contacts, and escape routes; "
    "moreover, translation and accommodation are provided as necessary for full comprehension by all crew "
    "members."
)
para(
    "II. Dynamic (\u201cRolling\u201d) Hazard Assessments: Workers and supervisors perform situational hazard "
    "assessments before any adjustment or movement of TTC devices, with particular focus on live traffic "
    "movement, weather, visibility, and prior incident locations; in addition, briefings at the end of each "
    "shift (\u201cdebriefings\u201d) discuss errors, near-misses, successful mitigations, and lessons for "
    "following shifts."
)
para(
    "III. Post-Incident/Adjustment Debriefs: After any significant task phase or incident, a structured review "
    "meeting documents what was done correctly and incorrectly and updates safety plans accordingly; this is "
    "especially important for high-risk shift transitions in multi-day or long-duration projects to ensure "
    "continuity in hazard awareness and TTC strategy (**Gambatese & Moeung, 2022**; Work Zone Safety "
    "Consortium, 2014)."
)
para(
    "IV. Communication Noteworthy Practices: Use of radios, hand signals, and spotters is critical for "
    "maintaining real-time communication during device movement/removal, especially in low-visibility or night "
    "operations; at no point should all workers face away from traffic, as at least one worker or spotter must "
    "always maintain vigilance toward live traffic; finally, a clearly rehearsed \u201cbailout plan\u201d for "
    "errant vehicles is discussed in briefings, with all crew aware of safe refuge areas and escape routes "
    "(ATSSA, 2008; ANSI/ISEA, 2020)."
)

para(
    "**The companion Field Pocket Guide frames worker protection as a set of layered controls; no single "
    "device or instruction is sufficient by itself. Table 14A summarizes these core protection layers.**"
)
para("**Table 14A: Core Worker-Protection Layers for Professional TTC Practice**")
table(
    ["Protection layer", "Minimum expectation", "Field emphasis"],
    [
        ["High-visibility safety apparel (HVSA)", "Compliant high-visibility apparel matched to exposure.", "Upgrade conspicuity for high-speed, night, and complex work, per ANSI/ISEA 107-2020."],
        ["Protective vehicles", "Operational shadow-vehicle or TMA protection where exposure warrants.", "Positioning and discipline matter more than mere presence."],
        ["Communications", "Reliable radios and clear commands.", "Brief roles, escape routes, and stop-work triggers."],
        ["Lighting", "Night work illuminated without excessive glare.", "Drivers must see the TTC; workers must see the hazard field."],
    ],
    note="**(Source: TTC Device Management Field Pocket Guide, Table 8; ANSI/ISEA, 2020; ATSSA, 2008.)**",
    bold_all_new=True,
)

para(
    "Noteworthy Practices in Worker Protection During TTC Activities: Pre-task inventory and readiness require "
    "checks of inventory and equipment, with function tests before deployment, and any non-compliant HVSA or "
    "TMA vehicle is immediately removed from active service (ATSSA, 2008; Gambatese & Moeung, 2022). In "
    "addition, minimize direct exposure by placing, moving, or retrieving devices from the shoulder or behind "
    "barriers to reduce time in live lanes, and note that automated or semi-automated equipment (e.g., cone "
    "placement machines) is increasingly specified. For night operations, provide proper supplemental lighting "
    "during installations/removals to meet reflectivity and visibility standards, and spotters become "
    "mandatory in these conditions (ANSI/ISEA, 2020). Finally, when possible, coordinate with law enforcement "
    "to provide speed monitoring and additional visibility for TTC deployment/removal operations (**ATSSA, "
    "2008**)."
)
para("Table 15: Typical Longitudinal Buffer Spaces and TMA Distances")
table(
    ["Speed Limit (mph)", "Buffer Space (ft)", "TMA Roll-Ahead (min, ft)"],
    [["35", "250", "74\u2013100"], ["55", "495", "150+"], ["65", "645", "150\u2013200"]],
)
para(
    "Empirically Derived Risk Factors and Noteworthy Countermeasures: First, Fact: Most work zone crashes "
    "involving workers occur during setting up, modifying, or removing TTC devices, not while the work zone is "
    "fully established. Moreover, Empirical Evidence: Risk increases acutely for workers during transitions, "
    "especially if high-visibility or attenuator standards are not meticulously followed. Accordingly, "
    "Actionable Response: Adherence to ANSI/ISEA 107-2020 standards, **MASH-compliant** TMA positioning, and "
    "structured, shift-based briefing protocols is non-negotiable for any reputable TTC operation."
)

para(
    "**Summary Field Recommendations. The lifecycle can be condensed into six field rules drawn from the "
    "companion Field Pocket Guide: (1) Plan\u2014start with MUTCD Part 6, agency standards, and the approved "
    "TCP; (2) Prepare\u2014verify device condition, visibility, message order, PPE, communications, and "
    "protection; (3) Place\u2014install advance warning, transition, activity, and termination elements in a "
    "clear sequence; (4) Adjust\u2014update protection when drivers appear confused, queues form, visibility "
    "drops, or workers are exposed; (5) Remove\u2014remove in reverse order, keeping the highest-value "
    "protection in place until the end; and (6) Verify\u2014drive through, correct residual hazards, document "
    "deficiencies, and share lessons learned. Safety must not be compromised in any work-duration category; "
    "simplified control must be balanced with other effective protective measures (TTC Device Management Field "
    "Pocket Guide).**"
)

# =========================================================================
# REFERENCES (revised, de-duplicated, corrected)
# =========================================================================
doc.add_page_break()
h("References", 1)
para("**(Revised reference list: duplicates removed, missing citations added, dates and publishers corrected. Newly added or corrected entries are shown in bold. See the List of Changes for details.)**")

refs = [
    "American Traffic Safety Services Association. (2008). Field guide on installation and removal of temporary traffic control for safe maintenance and work zone operations. ATSSA.",
    "**American Traffic Safety Services Association. (2025). Guidance for temporary traffic control in short-duration, short-term stationary, and mobile operations. ATSSA.**",
    "ANSI/ISEA. (2020). American national standard for high-visibility safety apparel and accessories (ANSI/ISEA 107-2020). American National Standards Institute / International Safety Equipment Association.",
    "**AASHTO. (2016). Manual for assessing safety hardware (MASH) (2nd ed.). American Association of State Highway and Transportation Officials.**",
    "Blincoe, L. J., Miller, T. R., Zaloshnja, E., & Lawrence, B. A. (2015). The economic and societal impact of motor vehicle crashes, 2010 (Revised May 2015) (Report No. DOT HS 812 013). National Highway Traffic Safety Administration. https://crashstats.nhtsa.dot.gov/Api/Public/ViewPublication/812013",
    "Bureau of Labor Statistics. (2020a). Census of fatal occupational injuries (CFOI) \u2014 current and revised data [2018 reference year]. U.S. Department of Labor. https://www.bls.gov/iif/oshcfoi1.htm",
    "Bureau of Labor Statistics. (2020b). Survey of occupational injuries and illnesses data [2018 reference year]. U.S. Department of Labor.",
    "Code of Federal Regulations. (2025). Title 23, Part 630, Subpart J \u2014 Work zone safety and mobility (23 CFR 630.1012). https://www.ecfr.gov/current/title-23/chapter-I/subchapter-G/part-630/subpart-J",
    "Connecticut Department of Transportation. (2017). Connecticut Department of Transportation smart work zones guide. Hartford, CT.",
    "**Federal Highway Administration. (2020). Commercial motor vehicle safety in work zones targeted action plan (Report No. FHWA-HOP-20-027). U.S. Department of Transportation.**",
    "Federal Highway Administration. (2021). Work zone intelligent transportation systems \u2014 technology supplement (Report No. FHWA-HOP-21-021). U.S. Department of Transportation. [Prepared by Schroeder, B., Warchol, S., Laffey, S., Grosso, R., Rowe, G., Pate, A., Boyapati, R., & Sanchez-Badillo, A.]",
    "**Federal Highway Administration. (2023). Manual on uniform traffic control devices for streets and highways (11th ed.). U.S. Department of Transportation. https://mutcd.fhwa.dot.gov**",
    "Finley, M. D., Theiss, L., Ullman, G. L., Pickens, A., Benden, M., & Jenkins, J. (2017). Evaluation of safety practices for short duration work zones (Report No. FHWA/OH-2017-29). Ohio Department of Transportation.",
    "Frydenberg, S., Aylward, K., Nordby, K., & Eikenes, J. O. H. (2021). Development of an augmented reality concept for icebreaker assistance and convoy operations. Journal of Marine Science and Engineering, 9(9), 996. https://doi.org/10.3390/jmse9090996",
    "Gambatese, J. A., & Moeung, S. R. (2022). Best practices for work zone safety during traffic control placement, removal, and modification \u2014 Phase I (SPR 839). Oregon State University / Oregon Department of Transportation / FHWA.",
    "**Gambatese, J. A., Moeung, S. R., Lee, W.-H., & Dai, Q. (2024). Best practices for work zone safety during traffic control placement, removal, and modification \u2014 Phase II (Report No. FHWA-OR-RD-24-04). Oregon State University / Oregon Department of Transportation / FHWA.**",
    "Illinois Department of Transportation. (2018). Work zone safety guidelines. Springfield, IL.",
    "Kansas Department of Transportation. (2021). Traffic control devices manual. Topeka, KS.",
    "Marjovi, A., Vasic, M., Lemaitre, J., & Martinoli, A. (2015). Distributed graph-based convoy control for networked intelligent vehicles. 2015 IEEE Intelligent Vehicles Symposium (IV). https://doi.org/10.1109/ivs.2015.7225676",
    "Massachusetts Department of Transportation. (2017). Work zone traffic control guidelines. Boston, MA.",
    "**National Work Zone Safety Information Clearinghouse. (2019). Work zone fatal crashes and fatalities [Data set]. https://workzonesafety.org/work-zone-data/work-zone-fatal-crashes-and-fatalities/**",
    "National Safety Council. (2024). Work zone safety training. https://www.nsc.org/safety-training/workplace/work-zone/work-zone-safety",
    "Oregon Department of Transportation. (2011). Oregon temporary traffic control handbook (OTTCH). https://www.oregon.gov/odot/Engineering/Docs/Traffic-Eng/OTTCH-v2011.pdf",
    "**Pennsylvania Department of Transportation. (2020). Work zone traffic control guidelines. Harrisburg, PA.**",
    "South Carolina Department of Transportation. (2019). Work zone traffic control procedure and guidelines for SCDOT maintenance activity. Columbia, SC.",
    "Texas A&M Engineering Extension Service. (2025). Work zone traffic control training. https://teex.org/class/hws002/",
    "Texas Department of Transportation. (2018). Smart work zone guidelines: Design guidelines for deployment of work zone intelligent transportation systems (ITS). Austin, TX.",
    "Theiss, L., & Ullman, G. L. (2017). Automated placement and retrieval of traffic cones (Report No. CDOT-2017-07). Colorado Department of Transportation; Texas A&M Transportation Institute.",
    "Theiss, L., Ullman, G. L., & Jackels, J. (2023). Safe and effective temporary traffic control for mobile operations on two-lane roadways. Texas A&M Transportation Institute.",
    "Theiss, L., Ullman, G. L., & Jackels, J. (2024). Improved temporary traffic control guidance for mobile operations on two-lane roadways. Texas A&M Transportation Institute.",
    "Ullman, G. L., Miller, M., Albert, D., Le, M., Crawford, J., & Brydia, R. (2024). Improving smart work zone deployments in Texas (Report No. FHWA/TX-24/0-7118-R1). Texas A&M Transportation Institute / Texas Department of Transportation.",
    "U.S. Department of Transportation. (2019). U.S. transportation secretary Elaine L. Chao announces further decreases in roadway fatalities. National Highway Traffic Safety Administration. https://www.nhtsa.gov",
    "**Utility Work Zone Traffic Control Guidelines. (2008). Wayne State University & Bradley University, FHWA Work Zone Safety Grant Program.**",
    "Virginia Department of Transportation. (2019). Virginia work area protection manual (2011 ed. with Revision 2). Richmond, VA.",
    "Virginia Department of Transportation. (2025). Intermediate work zone traffic control training and certification. https://catalog.reynolds.edu/preview_program.php?catoid=8&poid=1621",
    "Work Zone Safety Consortium. (2014). Guidance on safe temporary traffic control for short-duration and mobile operations. Federal Highway Administration.",
]
for i, r in enumerate(refs, 1):
    para(f"{i}. {r}", space_after=4)

# =========================================================================
# LIST OF CHANGES MADE
# =========================================================================
doc.add_page_break()
h("List of Changes Made", 1)
para(
    "Every substantive edit is itemized below. In the body of the document, all inserted or corrected text is "
    "shown in BOLD. Changes fall into five groups: (A) factual/statistical corrections, (B) standards "
    "updates, (C) reference-list corrections, (D) figure/picture source attributions, and (E) checklist and "
    "content additions from the companion Field Pocket Guide."
)

h("A. Factual and statistical corrections", 2)
A = [
    "Introduction: \u201cnearly 1,000 worker fatalities\u201d changed to the exact verified figure of 1,008 construction worker fatalities in 2018 (BLS CFOI 2018, Table 4). The phrase \u201cmore than three deaths per workday\u201d was changed to \u201cnearly three deaths every day\u201d (1,008 \u00f7 365 \u2248 2.8/day) to make the statement precisely verifiable.",
    "Introduction: Added the verified figure that the transportation and warehousing sector recorded 874 fatalities in 2018 (second-highest after construction), confirming the existing \u201csecond-highest\u201d statement. Fatality rates of 9.5 (construction) and 14.0 (transportation and warehousing) per 100,000 FTE, and 23.4 for agriculture/forestry/fishing/hunting, were all verified against BLS CFOI 2018 and retained.",
    "Introduction: Clarified that the construction nonfatal rate of 3.0 and transportation/warehousing rate of 4.5 exceed the private-industry average of 2.8 (BLS SOII 2018); \u201coverall industry average\u201d corrected to \u201cprivate-industry average\u201d for accuracy.",
    "Introduction: Corrected the figure \u201c238 crashes involving large trucks or buses\u201d to 215 fatal work zone crashes involving large trucks or buses (commercial motor vehicles)\u2014about 32% of all fatal work zone crashes in 2018 (FHWA, 2020, FHWA-HOP-20-027). The original 238 figure could not be verified in any source; 215 is the published value.",
    "Introduction: Attributed the 755 work zone fatalities / 124 worker deaths to the National Work Zone Safety Information Clearinghouse (2019), which is the actual data source; the prior in-text \u201cWorkzonesafety.org, 2019\u201d citation was not in the reference list.",
    "Section 3(c): Removed the unverifiable specific claim that TMA/shadow-vehicle positioning \u201cwithin 100 ft of device installers\u201d is \u201cshown to cut worker strikes,\u201d replacing it with positioning \u201cwithin recommended buffer distances\u2026 intended to reduce worker strikes,\u201d which is defensible.",
    "Section 2(d)/IV: Replaced the unverifiable \u201cODOT (2020)\u2026 measurable declines in struck-by incidents\u201d statement (no such 2020 ODOT publication exists in the literature reviewed) with a verifiable statement that field evaluations associate automation with reduced worker exposure and lower struck-by risk (Theiss & Ullman, 2017; Gambatese & Moeung, 2022).",
    "Section 5(d) Benefits: Softened the precise \u201creduce rear-end and secondary crashes by 40\u201355%\u201d claim, which could not be verified to the cited FHWA source, to a verifiable qualitative statement that queue-warning deployments report reductions in rear-end and secondary crashes (FHWA, 2021).",
    "Section 8: Reworded \u201cMost work zone crashes occur during setting up\u2026\u201d to \u201cMost work zone crashes involving workers occur during setting up\u2026\u201d (the original overstated the claim; most work zone fatalities are motorists).",
    "Table 1 (cone/drum spacing): Corrected the 70 mph taper spacing from 60 ft to 70 ft, consistent with the MUTCD rule that maximum channelizing-device spacing in a taper (in feet) equals the speed (in mph).",
]
for x in A:
    bullet(x)

h("B. Standards updates (replacing outdated references with current ones)", 2)
B = [
    "MUTCD edition/date corrected throughout: the document cited \u201cFHWA, 2025\u201d for the 11th-Edition MUTCD. The 11th Edition was published December 19, 2023 (effective January 18, 2024). All MUTCD citations now read \u201cMUTCD, 2023 / FHWA, 2023,\u201d and the reference entry was corrected to 2023.",
    "Truck-Mounted Attenuator (TMA) crashworthiness standard updated from NCHRP Report 350 (and the erroneous \u201cNCHRP 488\u201d) to AASHTO MASH (2016) in Section 5(c), Section 8(b), and Table 10. Added the AASHTO/FHWA implementation rule: TMAs manufactured after December 31, 2019 must meet MASH, while NCHRP 350 units made on/before that date may remain in service. A new AASHTO (2016) MASH reference was added.",
    "Added ATSSA (2025) \u201cGuidance for Temporary Traffic Control in Short-duration, Short-term Stationary, and Mobile Operations\u201d as a current source supporting the mobile and short-duration content (updating older ATSSA 2008 reliance where appropriate).",
    "Core spacing, taper, and buffer tables (Tables 1, 2, 4, 5, 6, 7, 8) were cross-checked against the MUTCD 11th Edition and annotated as consistent with MUTCD (2023); ATSSA (2008) retained where it is the specific field-guide source.",
    "Replaced the misattributed citation \u201cNIOSH (2020)\u201d for high-visibility apparel with \u201cANSI/ISEA (2020),\u201d which is the correct publisher of ANSI/ISEA 107-2020. (NIOSH does not publish that standard; the duplicate/misdescribed NIOSH entry was removed from the reference list.)",
]
for x in B:
    bullet(x)

h("C. Reference-list corrections (in-text \u2194 list reconciliation, duplicates, accuracy)", 2)
C = [
    "Added missing reference cited in text: National Work Zone Safety Information Clearinghouse (2019) for the 2018 work zone fatality data.",
    "Reconciled the non-existent \u201cGambatese & Moeung, 2023\u201d in-text citations: corrected to \u201c2022\u201d (Phase I) and added the genuine Phase II report \u2014 Gambatese, Moeung, Lee & Dai (2024), FHWA-OR-RD-24-04 \u2014 to the reference list as a recent source.",
    "Removed duplicate reference: the FHWA (2021) Work Zone ITS Technology Supplement (FHWA-HOP-21-021) was listed twice; the two \u201cSchroeder et al. (2021)\u201d entries are the same report, so they were merged into the single FHWA (2021) entry (author credit noted).",
    "Removed duplicate reference: Theiss & Ullman (2017) \u201cAutomated placement and retrieval of traffic cones\u201d was listed twice; kept the complete entry (Colorado DOT Report CDOT-2017-07; verified).",
    "Consolidated the duplicate Texas smart-work-zone reference: \u201cTTI (2024)\u201d and \u201cUllman et al. (2024)\u201d are the same project report (Improving Smart Work Zone Deployments in Texas, 0-7118-R1); in-text citations were standardized to \u201cUllman et al., 2024,\u201d and the duplicate list entry was removed.",
    "Split a merged reference entry: the original item 16 contained both Kansas DOT (2021) and Pennsylvania DOT (2020) in one entry; these are now two separate, correctly formatted references.",
    "Consolidated duplicate 23 CFR 630 Subpart J entries (the \u201cCode of Federal Regulations, 2025\u201d and \u201cFHWA, 2025\u201d CFR items were the same regulation); kept a single eCFR entry.",
    "Retained Finley et al. (2017) \u2014 previously uncited \u2014 by citing it in the short-duration work content (Section 4d), where it is directly relevant.",
    "Verified Theiss & Ullman (2017) is Colorado DOT Report CDOT-2017-07 (July 2017); corrected the inconsistent dual listing.",
    "Reframed the Marjovi et al. (2015) and Frydenberg et al. (2021) citations: these are general convoy-control / augmented-reality studies (intelligent-vehicle and marine icebreaker contexts), not work-zone studies. The work-zone convoy-management claim is now attributed to Theiss et al. (2023, 2024), and Marjovi/Frydenberg are cited only as broader supporting research, accurately described.",
]
for x in C:
    bullet(x)

h("D. Figure / picture source attributions added", 2)
D = [
    "Figure 4 (Component Parts of a TTC Zone): attributed to the MUTCD, 11th Edition, Part 6 (FHWA, 2023), which is the standard source of this layout.",
    "Figure 5 (Device Placement on Curves and Crests): annotated as illustrating MUTCD sight-distance principles (FHWA, 2023).",
    "Figure 6 (Rolling Roadblock with TMA): annotated as a concept after Theiss et al. (2023) and Gambatese & Moeung (2022).",
    "Figures 1, 2, 3, 7, 8, 9 are author-generated process/flow diagrams; no external source was asserted because none applies (avoiding fabricated citations). Figure 10 already carried a (MUTCD) attribution, retained.",
]
for x in D:
    bullet(x)

h("E. Checklist and content additions from the Field Pocket Guide", 2)
E = [
    "Section 3(a): Added Table 3A \u2014 Recommended Advance-Preparation (Pre-Installation) Readiness Checklist (10 checkbox items) from the Field Pocket Guide, with an introductory sentence.",
    "Section 4(b): Added field methods for approximating distances (roller tape, odometer, reference-location signs, station markings, skip-line, pacing) and the drive-through verification reminder.",
    "Section 4(c): Added Table 9A \u2014 Step-by-Step Installation of a Stationary Lane Closure (10-step field checklist) with an introduction and the protective-vehicle reminder.",
    "Section 4(d): Expanded the mobile/short-duration content with the Field Pocket Guide distinctions among short-term stationary (>1 hr), short-duration (\u22641 hr), and mobile operations, and the principle that \u201csimplified must never mean unprotected\u201d; added Finley et al. (2017) and ATSSA (2025) support.",
    "Section 7(b): Added Table 11A \u2014 Decisions for Small-City Contexts (short blocks, pedestrian/bicycle activity, constrained streets, local access/community impacts).",
    "Section 7(a): Added the Field Pocket Guide rules to open a detour only after the full route is signed/visible/understandable, and to confirm that normal route signs/markings are visible and non-conflicting after takedown.",
    "Section 7(e): Added Table 13A \u2014 Utility Work TTC Decisions (short-duration, lane closure, mobile/recurring, longer-duration/higher-risk) and the constrained-right-of-way and \u201crecheck after each move\u201d guidance; added the Utility Work Zone Traffic Control Guidelines (2008) reference.",
    "Section 8: Added Table 14A \u2014 Core Worker-Protection Layers (HVSA, protective vehicles, communications, lighting) and a Summary Field Recommendations paragraph (six field rules: Plan, Prepare, Place, Adjust, Remove, Verify).",
]
for x in E:
    bullet(x)

para(
    "**Note on scope:** This revision preserves the author\u2019s original structure, section order, and table "
    "numbering. Original figures (diagrams) are represented here by their captions with added source notes, "
    "since the underlying images reside in the original Word file; the bold caption notes can be copied "
    "directly into that file. All bolded text marks an addition or correction for easy review."
)

# =========================================================================
import os
out = os.path.join(os.path.dirname(os.path.abspath(__file__)), "TTCD_Management_Synthesis_REVISED.docx")
doc.save(out)
print("Saved:", out)

